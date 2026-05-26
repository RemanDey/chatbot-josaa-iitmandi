#!/usr/bin/env python3
"""
TASK 1 — Ingest + Normalize Docs (app/ingest.py)

Walks a raw data folder recursively, parses PDFs/text/jsonl, normalizes text,
chunks using RecursiveCharacterTextSplitter, computes local BGE embeddings,
stores chunks in a local ChromaDB collection named "iitmandi", and builds a BM25 index.
"""

import argparse
import hashlib
import json
import logging
import os
import pickle
import re
from datetime import datetime, timezone
from pathlib import Path

import chromadb
import fitz  # PyMuPDF
from chromadb.utils import embedding_functions
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from tqdm import tqdm

# ── Setup Logging ───────────────────────────────────────────────
os.makedirs("logs", exist_ok=True)
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    handlers=[
        logging.FileHandler("logs/rag.log", encoding="utf-8"),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger("ingest")

# Suppress noisy HuggingFace & transformers warnings/logs
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="huggingface_hub")
warnings.filterwarnings("ignore", category=UserWarning, module="transformers")
warnings.filterwarnings("ignore", message=".*unauthenticated requests to the HF Hub.*")
warnings.filterwarnings("ignore", message=".*unexpected keys.*")
warnings.filterwarnings("ignore", message=".*embeddings.position_ids.*")

logging.getLogger("huggingface_hub").setLevel(logging.ERROR)
logging.getLogger("transformers").setLevel(logging.ERROR)

# ── Constants ───────────────────────────────────────────────────
EMBEDDING_MODEL_NAME = "BAAI/bge-base-en-v1.5"
CHROMA_DB_DIR = "data/chroma_db"
HASHES_FILE = os.path.join(CHROMA_DB_DIR, "hashes.txt")
BM25_PICKLE_FILE = "data/bm25_corpus.pkl"
COLLECTION_NAME = "iitmandi"


def get_sha256(text: str) -> str:
    """Compute sha256 hash of a string."""
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def normalize_text(text: str) -> str:
    """
    Normalize text:
    - Strip multiple newlines (replace 3+ newlines with 2)
    - Strip leading/trailing whitespace per line
    - Remove lines that are only punctuation or single characters
    """
    # Replace 3+ newlines with 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    
    # Process line by line
    cleaned_lines = []
    for line in text.splitlines():
        line_stripped = line.strip()
        if not line_stripped:
            cleaned_lines.append("")
            continue
        # Check if line is only punctuation or single characters
        if len(line_stripped) <= 1 or re.match(r"^[^\w\s]+$", line_stripped):
            continue
        cleaned_lines.append(line_stripped)
        
    return "\n".join(cleaned_lines)


def infer_type_from_path(file_path: Path) -> str:
    """
    Infer document type based on parent folder name.
    Folder containing 'chat' -> 'chat', 'faq' -> 'faq', 'web' -> 'web', else 'doc'.
    """
    path_parts = [p.lower() for p in file_path.parts]
    if "chat" in path_parts:
        return "chat"
    elif "faq" in path_parts:
        return "faq"
    elif "web" in path_parts:
        return "web"
    return "doc"


def extract_pdf_text(file_path: Path) -> str:
    """Extract text from PDF using PyMuPDF."""
    text_content = []
    try:
        with fitz.open(str(file_path)) as doc:
            for page in doc:
                text_content.append(page.get_text())
    except Exception as e:
        logger.error("Error reading PDF %s: %s", file_path.name, e)
    return "\n".join(text_content)


def extract_docx_text(file_path: Path) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
        doc = docx.Document(str(file_path))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return "\n".join(full_text)
    except Exception as e:
        logger.error("Error reading DOCX %s: %s", file_path.name, e)
        return ""


def extract_jsonl_text(file_path: Path) -> list[tuple[str, dict]]:
    """
    Parse jsonl file where each line is {"instruction": str, "response": str}
    Returns list of (text_content, metadata) pairs.
    """
    parsed_docs = []
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            for i, line in enumerate(f, 1):
                line = line.strip()
                if not line:
                    continue
                try:
                    data = json.loads(line)
                    instruction = data.get("instruction", "").strip()
                    response = data.get("response", data.get("output", "")).strip()
                    if not instruction or not response:
                        continue
                    combined_text = f"Q: {instruction}\nA: {response}"
                    meta = {
                        "qa_index": i,
                        "category": data.get("category", "general")
                    }
                    parsed_docs.append((combined_text, meta))
                except Exception as e:
                    logger.warning("Malformed jsonl line at %s:%d: %s", file_path.name, i, e)
    except Exception as e:
        logger.error("Error reading JSONL %s: %s", file_path.name, e)
    return parsed_docs


def extract_markdown_frontmatter(text: str) -> tuple[str, dict]:
    """
    Extract YAML frontmatter from markdown text.
    Returns (cleaned_text, frontmatter_dict).
    """
    frontmatter = {}
    cleaned_text = text
    match = re.match(r"^---\s*\n(.*?)\n---\s*\n", text, re.DOTALL)
    if match:
        yaml_text = match.group(1)
        cleaned_text = text[match.end():]
        try:
            import yaml
            parsed = yaml.safe_load(yaml_text)
            if isinstance(parsed, dict):
                frontmatter = parsed
        except Exception as e:
            logger.warning("Failed to parse YAML frontmatter: %s. Using regex fallback.", e)
            for line in yaml_text.splitlines():
                line = line.strip()
                if not line or line.startswith("#"):
                    continue
                if ":" in line:
                    parts = line.split(":", 1)
                    k = parts[0].strip()
                    v = parts[1].strip()
                    if v.startswith("[") and v.endswith("]"):
                        try:
                            frontmatter[k] = json.loads(v)
                        except:
                            frontmatter[k] = v
                    else:
                        frontmatter[k] = v
    
    # Process lists (tags, aliases, related) to serialize them as JSON strings for ChromaDB
    flat_meta = {}
    for k, v in frontmatter.items():
        if isinstance(v, list):
            flat_meta[k] = json.dumps(v)
        else:
            flat_meta[k] = v
            
    return cleaned_text, flat_meta


def main():
    parser = argparse.ArgumentParser(description="Ingest and Normalize Documents")
    parser.path = parser.add_argument(
        "--folder", type=str, required=True, help="Path to raw data folder"
    )
    args = parser.parse_args()
    
    raw_folder = Path(args.folder)
    if not raw_folder.exists():
        logger.error("Raw folder does not exist: %s", raw_folder)
        return
        
    os.makedirs(CHROMA_DB_DIR, exist_ok=True)
    os.makedirs(os.path.dirname(BM25_PICKLE_FILE), exist_ok=True)
    
    # ── Load existing document hashes ─────────────────────────────
    existing_hashes = set()
    if os.path.exists(HASHES_FILE):
        with open(HASHES_FILE, "r", encoding="utf-8") as f:
            existing_hashes = {line.strip() for line in f if line.strip()}
            
    logger.info("Loaded %d existing document hashes.", len(existing_hashes))
    
    # ── Walk raw folder recursively ──────────────────────────────
    logger.info("Scanning folder: %s", raw_folder)
    all_files = []
    for root, _, files in os.walk(raw_folder):
        for file in files:
            all_files.append(Path(root) / file)
            
    logger.info("Found %d total files.", len(all_files))
    
    # We will accumulate document chunks, metadatas, and hashes to write
    new_chunks = []
    new_hashes_to_write = []
    
    # Initialize splitter with increased size to prevent overchunking and table splitting
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=200,
        separators=["\n\n", "\n", ". ", " "]
    )
    
    for file_path in tqdm(all_files, desc="Parsing files"):
        suffix = file_path.suffix.lower()
        if suffix not in [".pdf", ".txt", ".md", ".jsonl", ".docx"]:
            continue
            
        inferred_type = infer_type_from_path(file_path)
        ingested_at = datetime.now(timezone.utc).isoformat()
        
        # 1. Process files
        if suffix == ".pdf":
            raw_text = extract_pdf_text(file_path)
            documents_to_process = [(raw_text, {})]
        elif suffix == ".docx":
            raw_text = extract_docx_text(file_path)
            documents_to_process = [(raw_text, {})]
        elif suffix in [".txt", ".md"]:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    raw_text = f.read()
                if suffix == ".md":
                    cleaned_text, fm_meta = extract_markdown_frontmatter(raw_text)
                    documents_to_process = [(cleaned_text, fm_meta)]
                else:
                    documents_to_process = [(raw_text, {})]
            except Exception as e:
                logger.error("Error reading %s: %s", file_path.name, e)
                continue
        elif suffix == ".jsonl":
            documents_to_process = extract_jsonl_text(file_path)
            
        # 2. Process extracted documents
        for text, extra_meta in documents_to_process:
            text = text.strip()
            if not text or len(text) < 10:
                continue
                
            # Deduplicate by hashing full text
            doc_hash = get_sha256(text)
            if doc_hash in existing_hashes:
                continue
                
            # Normalize text
            normalized = normalize_text(text)
            if not normalized or len(normalized) < 10:
                continue
                
            # Metadata structure
            doc_metadata = {
                "source": file_path.name,
                "type": inferred_type,
                "ingested_at": ingested_at,
                **extra_meta
            }
            
            # Chunking with Title-Aware / Semantic prepending
            chunks = splitter.split_text(normalized)
            for idx, chunk_text in enumerate(chunks):
                chunk_text_stripped = chunk_text.strip()
                if not chunk_text_stripped:
                    continue
                
                # Prepend the document title to give the chunk semantic context
                doc_title = Path(file_path).stem.replace("_", " ").title()
                enhanced_text = f"Title: {doc_title}\n\n{chunk_text_stripped}"
                
                new_chunks.append({
                    "text": enhanced_text,
                    "metadata": {**doc_metadata, "chunk_index": idx}
                })
                
            existing_hashes.add(doc_hash)
            new_hashes_to_write.append(doc_hash)
            
    if not new_chunks:
        logger.info("No new documents to ingest. All already indexed or empty.")
        return
        
    logger.info("Prepared %d new chunks from new documents.", len(new_chunks))
    
    # ── Initialize local embedding model ──────────────────────────
    logger.info("Loading embedding model: %s", EMBEDDING_MODEL_NAME)
    embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
    
    # Compute embeddings in batches of 32
    logger.info("Computing embeddings (batch_size=32)...")
    chunk_texts = [c["text"] for c in new_chunks]
    embeddings = embedding_model.encode(
        chunk_texts,
        batch_size=32,
        show_progress_bar=True,
        normalize_embeddings=True
    )
    
    # ── Initialize ChromaDB ───────────────────────────────────────
    logger.info("Initializing ChromaDB collection '%s' at %s", COLLECTION_NAME, CHROMA_DB_DIR)
    chroma_client = chromadb.PersistentClient(path=CHROMA_DB_DIR)
    
    collection = chroma_client.get_or_create_collection(
        name=COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"}
    )
    
    # Ingest into ChromaDB in batches of 100
    batch_size = 100
    total_ingested = 0
    seen_ids = set()
    
    for i in tqdm(range(0, len(new_chunks), batch_size), desc="Ingesting to ChromaDB"):
        batch = new_chunks[i : i + batch_size]
        batch_embeddings = embeddings[i : i + batch_size]
        
        ids = []
        texts = []
        metadatas = []
        batch_embs = []
        
        for idx, c in enumerate(batch):
            # Compute ID as sha256(chunk_text)[:16]
            chunk_id = get_sha256(c["text"])[:16]
            if chunk_id in seen_ids:
                continue
            seen_ids.add(chunk_id)
            ids.append(chunk_id)
            texts.append(c["text"])
            
            # ChromaDB metadatas must be flat and simple scalar values (no nested dicts)
            flat_meta = {}
            for k, v in c["metadata"].items():
                if isinstance(v, (str, int, float, bool)):
                    flat_meta[k] = v
                else:
                    flat_meta[k] = str(v)
            metadatas.append(flat_meta)
            batch_embs.append(batch_embeddings[idx].tolist())
            
        if ids:
            collection.upsert(
                ids=ids,
                documents=texts,
                embeddings=batch_embs,
                metadatas=metadatas
            )
            total_ingested += len(ids)
            
    logger.info("Successfully ingested %d chunks to ChromaDB.", total_ingested)
    
    # Write new hashes to hashes.txt
    if new_hashes_to_write:
        with open(HASHES_FILE, "a", encoding="utf-8") as f:
            for h in new_hashes_to_write:
                f.write(h + "\n")
                
    # ── Build and Save BM25 Index ─────────────────────────────────
    logger.info("Building BM25 index...")
    # Load all existing chunks in collection to build a complete BM25 index
    all_data = collection.get(include=["documents", "metadatas"])
    all_documents = all_data.get("documents", [])
    all_metadatas = all_data.get("metadatas", [])
    
    bm25_corpus = []
    tokenized_corpus = []
    
    for doc, meta in zip(all_documents, all_metadatas):
        # Tokenize (lowercase, split on whitespace)
        tokens = doc.lower().split()
        bm25_corpus.append({"text": doc, "metadata": meta})
        tokenized_corpus.append(tokens)
        
    logger.info("Saving BM25 corpus (%d documents) to %s", len(bm25_corpus), BM25_PICKLE_FILE)
    with open(BM25_PICKLE_FILE, "wb") as f:
        pickle.dump({
            "corpus": bm25_corpus,
            "tokenized_corpus": tokenized_corpus
        }, f)
        
    logger.info("Ingestion process complete!")


if __name__ == "__main__":
    main()
